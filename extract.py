"""
Pull just the finished PROSE out of a .docx into a clean .txt, so the
detector scores real writing instead of headings, bullets and note-scaffolding.

Drops: headings, cover/TOC lines, footnotes (python-docx never reads those),
tables (also not in doc.paragraphs), bullet lines, and my note-labels
(NOTES / Verdict / Analysis: / Criticism / Mini-conclusion ...).

    python extract.py "in.docx" out.txt
"""

import re
import sys
import docx

MIN_WORDS = 25   # same floor as detect.py: shorter = fragment, not prose

# Lines that begin with any of these are note-scaffolding, not prose.
NOTE_STARTS = (
    "NOTES", "Verdict", "Analysis:", "Criticism", "Mini-conclusion",
    "Tool definition", "Why chosen", "Overall synthesis", "Force =",
    "POINTS", "Safest", "Moderate", "Research Question", "Word count",
    "Table of Contents",
)
# ponytail: prefix heuristic, not a parser. If a real sentence ever starts
# with one of these words it gets dropped too — check the .txt if a section vanishes.


def is_prose(p):
    text = p.text.strip()
    if len(text.split()) < MIN_WORDS:
        return False
    if p.style.name.startswith("Heading"):
        return False
    if text[0] in "•-*“\"[»":               # bullet, quoted snippet, or [SCAFFOLD]/» note marker
        return False
    if text.startswith(NOTE_STARTS):
        return False
    return True


def main():
    if len(sys.argv) != 3:
        sys.exit("usage: python extract.py in.docx out.txt")
    in_path, out_path = sys.argv[1], sys.argv[2]
    doc = docx.Document(in_path)
    prose = []
    for p in doc.paragraphs:
        # Bibliography is the last section; once its heading shows up, the rest
        # is citations, not prose. Stop here.
        if p.style.name.startswith("Heading") and p.text.strip().lower() in (
            "bibliography", "works cited", "references",
        ):
            break
        if is_prose(p):
            prose.append(p.text.strip())
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(prose))
    words = sum(len(p.split()) for p in prose)
    print(f"kept {len(prose)} paragraphs, {words} words -> {out_path}")


if __name__ == "__main__":
    main()
