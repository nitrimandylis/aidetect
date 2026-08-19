"""
Pull just the finished PROSE out of a .docx into a clean .txt, so the
detector scores real writing instead of headings, bullets and note-scaffolding.

Drops: headings, cover/TOC lines, footnotes (python-docx never reads those),
tables (also not in doc.paragraphs), bullet lines, and my note-labels
(NOTES / Verdict / Analysis: / Criticism / Mini-conclusion ...).

    aidetect extract "in.docx" out.txt
"""

import argparse

import docx

from .text import MIN_WORDS, is_end_heading, is_prose


def main(argv=None):
    ap = argparse.ArgumentParser(prog="aidetect extract",
                                 description="Extract finished prose from a .docx.")
    ap.add_argument("in_path", help="source .docx")
    ap.add_argument("out_path", help="destination .txt")
    args = ap.parse_args(argv)

    doc = docx.Document(args.in_path)
    prose = []
    for p in doc.paragraphs:
        # Bibliography is the last section; once its heading shows up, the rest
        # is citations, not prose. Stop here.
        if is_end_heading(p):
            break
        if is_prose(p, min_words=MIN_WORDS):
            prose.append(p.text.strip())
    with open(args.out_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(prose))
    words = sum(len(p.split()) for p in prose)
    print(f"kept {len(prose)} paragraphs, {words} words -> {args.out_path}")
    print("note: this is the detector's prose filter, not an IB count. "
          "Use `aidetect count` for that.")
