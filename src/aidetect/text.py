"""
Reading prose out of .docx and .txt files. No torch in here on purpose:
`aidetect count` uses this module and should not pay a model import to count words.

Two layers, deliberately separate:

  walk()      structure. What the IB excludes because of WHERE it sits in the
              document: the cover page, the contents page, headings, figure
              captions, tables, footnotes, and everything from the bibliography
              on. `count` and `extract` use this and nothing else.

  is_prose()  style. What the DETECTOR should not be handed: fragments under 25
              words, bullet and quote lines, list items, table rows, labelled
              criteria, and my own note-scaffolding. These are real words the
              examiner counts, so they never touch `count`.
"""

import re

MIN_WORDS = 25   # ponytail: skip fragments; short text scores as noise

# Lines that begin with any of these are note-scaffolding, not prose.
# ponytail: prefix heuristic, not a parser. If a real sentence ever starts
# with one of these words it gets dropped too — check the output if a section vanishes.
# Cover-page and contents lines are NOT listed here: walk() drops those by
# position, which is why "Word count:" and "Table of Contents" are gone.
NOTE_STARTS = (
    "NOTES", "Verdict", "Analysis:", "Criticism", "Mini-conclusion",
    "Tool definition", "Why chosen", "Overall synthesis", "Force =",
    "POINTS", "Safest", "Moderate",
)

# Headings that mean "the prose is over, the rest is citations".
END_HEADINGS = ("bibliography", "works cited", "references")

# Headings whose whole section is excluded, but which the prose resumes after.
SKIP_SECTIONS = ("table of contents", "contents")

# A figure caption, not prose. The colon is required on purpose: "Figure 4:
# Graph showing Sales Revenue" is a caption, but "Figure 4 shows revenue rising"
# is a sentence the examiner counts, and only the colon separates them.
CAPTION = re.compile(r"^(?:Figure|Table|Fig\.)\s*\d+\s*:", re.IGNORECASE)


def is_heading(p):
    return p.style.name.startswith("Heading") or p.style.name == "Title"


def heading_level(p):
    """1 for Heading 1 and Title, 2 for Heading 2, and so on. 0 if not a heading.

    Title sits above Heading 1 in Word, but a Title is the essay's own name on
    the cover page, so treating it as level 1 is close enough and keeps the
    report from growing a level nobody uses.
    """
    if not is_heading(p):
        return 0
    tail = p.style.name.removeprefix("Heading ")
    if tail.isdigit():
        return int(tail)
    return 1


def is_end_heading(p):
    return is_heading(p) and p.text.strip().lower() in END_HEADINGS


def has_headings(path):
    """True if the document has at least one heading to anchor the cover-page
    rule to. A draft with none is counted whole, with a warning."""
    import docx
    return any(is_heading(p) for p in docx.Document(path).paragraphs)


def walk(path):
    """Yield (level, section_title, text) for every countable paragraph.

    `level` and `section_title` describe the heading the paragraph sits under.
    Structure exclusions only — see the module docstring for the split.

    Tables and footnotes need no handling here: doc.paragraphs excludes table
    cells, and python-docx never exposes footnotes at all. The IB excludes both.
    """
    import docx  # python-docx, only needed for Word files

    doc = docx.Document(path)
    # No heading yet means we are still on the cover page, which never counts.
    # A headingless draft would therefore count as zero, so it is counted whole
    # instead: see has_headings() and the caller's warning.
    whole_document = not any(is_heading(p) for p in doc.paragraphs)
    level = 1 if whole_document else 0
    title = "(no headings)" if whole_document else None
    skipping = False

    for p in doc.paragraphs:
        if is_end_heading(p):
            return                      # bibliography onward is citations
        if is_heading(p):
            heading_text = p.text.strip()
            if not heading_text:
                continue                # an empty heading is a page break
            level = heading_level(p)
            title = heading_text
            skipping = heading_text.lower() in SKIP_SECTIONS
            yield (level, title, None)  # announce the section, even if empty
            continue
        if title is None or skipping:
            continue                    # cover page, or the contents page
        text = p.text.strip()
        if not text or CAPTION.match(text):
            continue
        yield (level, title, text)


# List items, table rows and labelled criteria. Not prose, even though they are
# words on a line, and `count` still counts them because the examiner does.
#
# This exists because of a measured failure. Segment mode drops the 25-word
# floor on purpose, so that short connective sentences get scored inside a
# window instead of being skipped. The side effect was that a CS IA's
# success-criteria table ("SC12  Highlight the active filter") and its numbered
# design breakdown ("1.", "2.") were handed to the detector as prose, and
# between them they drove 16% of that draft into the flagged bucket. That is
# the detector measuring formatting, not writing.
#
# Deliberately NOT here: single-letter items like "a)" and roman numerals like
# "V.". Both are indistinguishable from an initial at the start of a sentence,
# and dropping a paragraph that opens "T. S. Eliot wrote..." would be a worse
# error than keeping a stray list item.
LIST_MARKERS = re.compile(
    r"^("
    r"[•▪◦·*+–—-](\s|$)"     # bullet of some kind
    r"|\d+[.)](\s|$)"          # 1.  or  1)  or a bare "2." on its own line
    r"|[A-Z]{1,5}\d+[\s:]"     # labelled criterion: SC12, FR3, NFR10
    r")"
)


def is_list_item(text):
    """True for a bullet, a numbered item, a labelled criterion or a table row.

    A tab means columns: python-docx keeps real table cells out of
    doc.paragraphs, but a table typed as tab-separated text is still a table.
    """
    stripped = text.strip()
    if not stripped:
        return False
    if "\t" in stripped:
        return True
    return LIST_MARKERS.match(stripped) is not None


def is_prose(text, min_words=MIN_WORDS):
    """True if this paragraph is finished prose, not scaffolding.

    Style filter for the detector only. `count` must not use it: a 6-word
    sentence and a quoted passage are both words the examiner counts.
    """
    text = text.strip()
    if len(text.split()) < min_words:
        return False
    if text[0] in "•-*“\"[»":               # bullet, quoted snippet, or [SCAFFOLD]/» note marker
        return False
    if is_list_item(text):
        return False
    if text.startswith(NOTE_STARTS):
        return False
    return True


def read_paragraphs(path, min_words=MIN_WORDS):
    """Prose paragraphs for the detector, from a .docx or a .txt."""
    if path.lower().endswith(".docx"):
        chunks = [text for _level, _title, text in walk(path) if text is not None]
    else:
        with open(path, encoding="utf-8") as f:
            # blank line separates paragraphs in a plain-text file
            chunks = f.read().split("\n\n")
    return [c.strip() for c in chunks if is_prose(c, min_words)]


def bar(prob, width=20):
    filled = round(prob * width)
    return "#" * filled + "-" * (width - filled)
