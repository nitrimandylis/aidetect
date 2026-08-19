"""
IB word count for a draft.

Word's own count is wrong for IB: it counts headings, block quotes, tables,
footnotes and citations, none of which are assessed. This counts what the IB
counts, and splits the total by section so an over-long draft says WHERE.

Excluded, per the IB rules that are common to the EE and the subject IAs:
  - headings                  (skipped by style)
  - the bibliography onward   (stop at the Bibliography/Works Cited heading)
  - block quotes and bullets  (same first-character rule extract.py uses)
  - tables and footnotes      (python-docx never puts them in doc.paragraphs)
  - in-text citations         (parentheticals containing a year / ibid / et al)

    aidetect count draft.docx
    aidetect count draft.docx --limit 4000
    aidetect count draft.docx --limit 4000 --json
"""

import argparse
import json
import re

from .text import is_end_heading, is_heading, is_prose

# A parenthetical is a citation if it carries a 4-digit year, "ibid" or "et al".
# ponytail: narrow on purpose. "(the second of these)" is prose and stays counted;
# widening this to every parenthetical would silently eat hundreds of real words.
CITATION = re.compile(
    r"\((?:[^()]*(?:\b(?:1[6-9]|20)\d{2}\b|\bibid\b|\bet al\b)[^()]*)\)",
    re.IGNORECASE,
)


def count_words(text):
    """Words in one paragraph, with in-text citations removed.

    Tokens with no letter or digit are dropped: removing "(Smith, 2024)" from
    "...sharply (Smith, 2024)." strands the full stop as its own token, and a
    lone "." is not a word.
    """
    stripped = CITATION.sub(" ", text)
    return sum(1 for token in stripped.split() if any(c.isalnum() for c in token))


def count_docx(path):
    """Return (sections, total). sections is a list of (heading, words)."""
    import docx

    sections = [("(untitled)", 0)]
    for p in docx.Document(path).paragraphs:
        if is_end_heading(p):
            break
        if is_heading(p):
            title = p.text.strip()
            if title:
                sections.append((title, 0))
            continue
        # min_words=1: the 25-word floor is a detector heuristic. A one-line
        # paragraph is still words the examiner counts.
        if is_prose(p, min_words=1):
            title, words = sections[-1]
            sections[-1] = (title, words + count_words(p.text))
    sections = [s for s in sections if s[1] > 0]
    return sections, sum(w for _, w in sections)


def report(sections, total, limit):
    width = max((len(t) for t, _ in sections), default=10)
    for title, words in sections:
        share = words / total if total else 0
        print(f"  {title:<{width}}  {words:>5}  {'#' * round(share * 30)}")
    print("-" * (width + 40))
    if limit:
        over = total - limit
        verdict = f"{over:+d} over" if over > 0 else f"{-over} to spare"
        print(f"  {'TOTAL':<{width}}  {total:>5}  / {limit} limit, {verdict}")
    else:
        print(f"  {'TOTAL':<{width}}  {total:>5}")
    print("excludes headings, quotes, bullets, tables, footnotes, "
          "citations and everything from the bibliography on.")


def as_json(sections, total, limit):
    """The machine interface. Every key is always present; `null` means the
    question does not apply, not that it could not be answered. A draft with no
    prose is an empty list and exit 0, because parsing succeeded and the honest
    answer is zero."""
    return {
        "sections": [{"title": t, "words": w} for t, w in sections],
        "total": total,
        "limit": limit,
        "over": None if limit is None else total - limit,
    }


def main(argv=None):
    ap = argparse.ArgumentParser(prog="aidetect count",
                                 description="IB-rules word count for a draft.")
    ap.add_argument("path", help=".docx file to count")
    ap.add_argument("--limit", type=int, default=None,
                    help="word limit to measure against (e.g. 4000 for the EE)")
    ap.add_argument("--json", action="store_true",
                    help="emit one JSON object on stdout instead of the table")
    args = ap.parse_args(argv)

    if not args.path.lower().endswith(".docx"):
        ap.error("count needs a .docx (a .txt has no headings or styles to go on)")

    sections, total = count_docx(args.path)

    if args.json:
        print(json.dumps(as_json(sections, total, args.limit)))
        return
    if not sections:
        print("No prose found. Check the file is a draft and not an outline.")
        return
    report(sections, total, args.limit)
