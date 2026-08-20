"""
The count subcommand decides a number I hand to an examiner, so the rules it
applies are worth pinning down. Run with `python -m pytest tests/test_count.py`,
or just `python tests/test_count.py`.

The fixture is synthetic on purpose: my real drafts are coursework and do not
belong in a public package, and their numbers change every time I redraft.
"""

import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "src"))

from aidetect.count import count_docx, count_words, rolled_up  # noqa: E402


def test_citations_are_not_counted():
    # Harvard, with and without a page number, and the two Latin forms.
    assert count_words("The market grew sharply (Smith, 2024).") == 4
    assert count_words("The market grew sharply (Smith, 2024, p. 14).") == 4
    assert count_words("The market grew sharply (ibid.).") == 4
    assert count_words("The market grew sharply (Smith et al.).") == 4


def test_ordinary_parentheses_are_counted():
    # This is the whole reason the rule keys on a year rather than on brackets.
    assert count_words("The second option (the cheaper one) was chosen.") == 8


def test_year_in_running_prose_still_counts():
    # Only the parenthetical is stripped, not every year on the page.
    assert count_words("Revenue fell in 2024 across every region.") == 7


def test_count_never_imports_torch():
    """`aidetect count` must stay instant. A stray torch import in the chain
    from cli -> count -> text would add ~2s to counting words.

    Checked in a subprocess: asserting on sys.modules in-process would pass or
    fail depending on whether some other test file imported torch first.
    """
    import subprocess
    src = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "src"))
    env = dict(os.environ, PYTHONPATH=src)
    code = ("import sys, aidetect.cli, aidetect.count; "
            "sys.exit(1 if 'torch' in sys.modules else 0)")
    assert subprocess.run([sys.executable, "-c", code], env=env).returncode == 0


def build_fixture(path):
    """A draft shaped like my real ones: cover page, contents, nested headings,
    a caption, a table and a bibliography."""
    import docx
    doc = docx.Document()

    # Cover page: no heading above it, so none of this counts.
    doc.add_paragraph("Business Management")
    doc.add_paragraph("Word count: 1234/1800")

    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("Introduction ..................... 1")

    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This introduction runs to exactly seven words.")     # 7

    doc.add_heading("Analysis and Discussion", level=1)   # no words of its own
    doc.add_heading("Ethical Discussion", level=2)
    doc.add_paragraph("Revenue rose over the period (Smith, 2024).")        # 5
    doc.add_paragraph("“A block quote that must count in full.”")           # 8
    doc.add_paragraph("Figure 1: Graph showing revenue from 2021 to 2025.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Table cells must not count either."

    doc.add_heading("Bibliography", level=1)
    doc.add_paragraph("Smith, J. (2024). A Book That Must Not Count. Press.")
    doc.save(path)


def test_structure_exclusions():
    """Cover page, contents, captions, tables and bibliography all stay out;
    the block quote stays in."""
    path = "/tmp/aidetect-count-fixture.docx"
    build_fixture(path)
    sections, total = count_docx(path)
    assert sections == [
        (1, "Introduction", 7),
        (1, "Analysis and Discussion", 0),   # kept: it has children
        (2, "Ethical Discussion", 13),       # 5 + the 8-word block quote
    ], sections
    assert total == 20
    os.remove(path)


def test_own_words_sum_to_the_total():
    """The invariant the JSON contract rests on: no section double-counts its
    children, so a consumer can add the list up and get the reported total."""
    path = "/tmp/aidetect-count-sum.docx"
    build_fixture(path)
    sections, total = count_docx(path)
    assert sum(words for _level, _title, words in sections) == total
    os.remove(path)


def test_rollup_gives_a_parent_its_children():
    sections = [(1, "Intro", 7), (1, "Analysis", 0),
                (2, "Ethics", 13), (3, "Sub", 4), (1, "Conclusion", 5)]
    # Analysis owns Ethics and Sub and stops at Conclusion.
    assert rolled_up(sections) == [7, 17, 17, 4, 5]


def test_a_figure_sentence_is_not_a_caption():
    """"Figure 4:" is a caption. "Figure 4 shows" is assessed prose. The colon
    is the only thing telling them apart, so it is required."""
    import docx
    path = "/tmp/aidetect-count-caption.docx"
    doc = docx.Document()
    doc.add_heading("Analysis", level=1)
    doc.add_paragraph("Figure 4: Graph showing revenue.")            # caption
    doc.add_paragraph("Figure 4 shows that revenue grew sharply.")   # 7 words
    doc.save(path)
    sections, total = count_docx(path)
    assert sections == [(1, "Analysis", 7)], sections
    assert total == 7
    os.remove(path)


def test_a_draft_with_no_headings_is_counted_whole():
    """The cover-page rule is 'nothing counts until the first heading'. Applied
    blindly to a headingless rough draft that reports a confident zero, which is
    worse than counting the cover page."""
    import docx
    from aidetect.text import has_headings
    path = "/tmp/aidetect-count-noheadings.docx"
    doc = docx.Document()
    doc.add_paragraph("This rough draft has no headings at all in it.")   # 10
    doc.add_paragraph("It should still be counted rather than reported as zero.")  # 10
    doc.save(path)
    sections, total = count_docx(path)
    assert total == 20, sections
    assert sections == [(1, "(no headings)", 20)], sections
    assert has_headings(path) is False
    os.remove(path)


def run_cli(args):
    """Run the real entry point in a subprocess. Returns (rc, stdout, stderr).
    In-process would not catch the two things worth checking: the exit code and
    whether anything else leaked onto stdout."""
    import subprocess
    src = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "src"))
    env = dict(os.environ, PYTHONPATH=src)
    p = subprocess.run([sys.executable, "-c", "import sys; from aidetect.cli import main; sys.exit(main() or 0)",
                        *args], env=env, capture_output=True, text=True)
    return p.returncode, p.stdout, p.stderr


def test_json_is_the_only_thing_on_stdout():
    import json
    path = "/tmp/aidetect-json-fixture.docx"
    build_fixture(path)
    rc, out, _ = run_cli(["count", path, "--limit", "10", "--json"])
    assert rc == 0, rc
    payload = json.loads(out)          # fails if a banner or a warning got in
    assert payload == {
        "sections": [{"title": "Introduction", "words": 7, "level": 1},
                     {"title": "Analysis and Discussion", "words": 0, "level": 1},
                     {"title": "Ethical Discussion", "words": 13, "level": 2}],
        "total": 20,
        "limit": 10,
        "over": 10,
    }, payload
    os.remove(path)


def test_json_keys_stay_present_without_a_limit():
    # null means "does not apply"; the key is never dropped.
    import json
    path = "/tmp/aidetect-json-nolimit.docx"
    build_fixture(path)
    _, out, _ = run_cli(["count", path, "--json"])
    payload = json.loads(out)
    assert payload["limit"] is None and payload["over"] is None, payload
    os.remove(path)


def test_bad_flag_and_bad_input_exit_nonzero():
    # An unknown flag must be an error, not a silent shrug, and a .txt must not
    # be counted as if it were a draft.
    rc, _, _ = run_cli(["count", "whatever.docx", "--bogus"])
    assert rc != 0, "unknown flag exited 0"
    rc, _, err = run_cli(["count", "notes.txt"])
    assert rc != 0 and "docx" in err, (rc, err)


def test_extract_defaults_its_output_path():
    """`aidetect extract draft.docx` writes '<name> prose.txt' beside it, and
    the words in that file are the ones count counted (before citations)."""
    path = "/tmp/aidetect-extract-fixture.docx"
    out_path = "/tmp/aidetect-extract-fixture prose.txt"
    build_fixture(path)
    rc, out, err = run_cli(["extract", path])
    assert rc == 0, (rc, err)
    assert out_path in out, out
    with open(out_path, encoding="utf-8") as f:
        prose = f.read()
    assert "Business Management" not in prose      # cover page
    assert "Figure 1:" not in prose                # caption
    assert "Must Not Count" not in prose           # bibliography
    assert "block quote" in prose                  # quotes are assessed prose
    # 22 raw words: the 20 counted plus the two of "(Smith, 2024)".
    assert len(prose.split()) == 22, prose.split()
    os.remove(path)
    os.remove(out_path)


if __name__ == "__main__":
    for name, fn in sorted(globals().items()):
        if name.startswith("test_"):
            fn()
            print(f"ok  {name}")
    print("all count checks passed")
